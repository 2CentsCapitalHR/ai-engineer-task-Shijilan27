from pathlib import Path

from docx import Document


def make_aoa_example(path: Path) -> None:
    doc = Document()
    doc.add_heading("Articles of Association", level=1)
    doc.add_paragraph("This document sets out the governance of the Company.")
    doc.add_paragraph("Share capital and shareholder rights are addressed herein.")
    doc.add_paragraph("Execution: ")
    doc.save(path)


def make_application_example(path: Path) -> None:
    doc = Document()
    doc.add_heading("Incorporation Application Form", level=1)
    doc.add_paragraph("Applicant details and registered address.")
    doc.add_paragraph("The company seeks to be incorporated in Abu Dhabi Global Market (ADGM).")
    doc.add_paragraph("Execution: Signed by authorized signatory. Date: ")
    doc.save(path)


def main() -> None:
    out_dir = Path("examples")
    out_dir.mkdir(parents=True, exist_ok=True)
    make_aoa_example(out_dir / "AoA_example.docx")
    make_application_example(out_dir / "Incorporation_Application_example.docx")
    print(f"Examples saved to {out_dir.resolve()}")


if __name__ == "__main__":
    main()

