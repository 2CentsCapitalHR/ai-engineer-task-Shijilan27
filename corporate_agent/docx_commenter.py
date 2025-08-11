from typing import Dict, List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_comment(paragraph, author: str, text: str) -> None:
    try:
        # Create a comment range around the whole paragraph
        p = paragraph._p
        start = OxmlElement("w:commentRangeStart")
        end = OxmlElement("w:commentRangeEnd")
        # Use a simple incremental id based on existing comments
        root = p.getroottree()
        existing = root.xpath("//w:commentRangeStart", namespaces=p.nsmap)
        comment_id = str(len(existing) + 1)
        start.set(qn("w:id"), comment_id)
        end.set(qn("w:id"), comment_id)
        p.addprevious(start)
        p.addnext(end)

        # Ensure comments part exists
        doc = paragraph._element.getroottree().getroot().part
        comments_part = doc.comments_part
        if comments_part is None:
            comments_part = doc.add_comments_part()
        comments_el = comments_part.element

        # Build the <w:comment>
        comment = OxmlElement("w:comment")
        comment.set(qn("w:id"), comment_id)
        comment.set(qn("w:author"), author)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.text = text
        r_el.append(t_el)
        p_el.append(r_el)
        comment.append(p_el)
        comments_el.append(comment)
    except Exception:
        # Fallback: append a reviewer note run within the same paragraph
        paragraph.add_run(f"\n[Reviewer Note - {author}] {text}")


def build_reviewed_docx(doc: Document, issues: List[Dict]) -> Document:
    if not issues:
        return doc

    # Naive mapping: attach a comment for each issue to the first paragraph containing a related keyword
    for issue in issues:
        note = issue.get("issue", "Issue")
        section = issue.get("section", "Section")
        severity = issue.get("severity", "Medium")
        suggestion = issue.get("suggestion")
        refs = issue.get("references") or []
        ref_str = "; ".join([
            f"{r.get('title')} [{r.get('source')}] — {r.get('snippet','')[:80]}" for r in refs
        ])
        comment_text = f"{section} | {severity}: {note}"
        if suggestion:
            comment_text += f"\nSuggestion: {suggestion}"
        if ref_str:
            comment_text += f"\nRefs: {ref_str}"

        attached = False
        keywords = [k for k in [issue.get("section"), "adgm", "jurisdiction", "signature"] if k]
        for paragraph in doc.paragraphs:
            content = paragraph.text.lower()
            if any(k.lower() in content for k in keywords):
                _add_comment(paragraph, author="Corporate Agent", text=comment_text)
                attached = True
                break
        if not attached and doc.paragraphs:
            _add_comment(doc.paragraphs[0], author="Corporate Agent", text=comment_text)

    return doc

