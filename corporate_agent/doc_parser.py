from typing import List

from docx import Document


def load_docx_from_bytes(buffer: bytes) -> Document:
    from io import BytesIO

    return Document(BytesIO(buffer))


def extract_full_text(doc: Document) -> str:
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)
    return "\n".join(parts)

